from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt

# ایجاد یک سند جدید
doc = Document()
doc.add_heading('هوک‌های ری‌اکت و نکست‌جی‌اس', 0)

# افزودن مقدمه
doc.add_heading('مقدمه', level=1)
doc.add_paragraph(
    "ری‌اکت (React) و نکست‌جی‌اس (Next.js) از کتابخانه‌ها و فریمورک‌های محبوب جاوااسکریپت برای ساخت اپلیکیشن‌های تحت وب هستند. "
    "هوک‌ها (Hooks) در ری‌اکت امکانات قدرتمندی برای مدیریت وضعیت و چرخه زندگی کامپوننت‌ها ارائه می‌دهند. "
    "در اینجا به معرفی تمام هوک‌های اصلی ری‌اکت و نکست‌جی‌اس به همراه مثال و توضیحات می‌پردازیم."
)

# تعریف سبک Code
style = doc.styles.add_style('Code', 1)
style.font.name = 'Courier New'
style.font.size = Pt(10)

# هوک‌های ری‌اکت
doc.add_heading('هوک‌های ری‌اکت', level=1)

# useState
doc.add_heading('useState', level=2)
doc.add_paragraph(
    "برای مدیریت وضعیت محلی (local state) یک کامپوننت استفاده می‌شود. "
    "useState یک جفت مقدار برمی‌گرداند: مقدار کنونی وضعیت و تابعی برای به‌روزرسانی آن. "
    "می‌توان چندین useState برای مدیریت چندین وضعیت محلی در یک کامپوننت استفاده کرد."
)
doc.add_paragraph(
    """```javascript
import React, { useState } from 'react';

function Counter() {
  const [count, setCount] = useState(0);

  return (
    <div>
      <p>Count: {count}</p>
      <button onClick={() => setCount(count + 1)}>Increment</button>
    </div>
  );
}
```""",
    style='Code'
)

# useEffect
doc.add_heading('useEffect', level=2)
doc.add_paragraph(
    "برای مدیریت اثرات جانبی مانند درخواست‌های API، تایمرها و دسترسی به DOM استفاده می‌شود. "
    "useEffect تابعی است که بعد از رندر کردن کامپوننت اجرا می‌شود. "
    "می‌توان از آرایه‌ای به‌عنوان دومین آرگومان استفاده کرد تا مشخص شود چه زمانی این اثر جانبی باید اجرا شود (وابستگی‌ها)."
)
doc.add_paragraph(
    """```javascript
import React, { useEffect, useState } from 'react';

function DataFetcher() {
  const [data, setData] = useState(null);

  useEffect(() => {
    fetch('https://api.example.com/data')
      .then(response => response.json())
      .then(data => setData(data));
  }, []); // این آرایه‌ی خالی یعنی فقط یک بار بعد از اولین رندر اجرا شود.

  return <div>{data ? data : 'Loading...'}</div>;
}
```""",
    style='Code'
)

# useContext
doc.add_heading('useContext', level=2)
doc.add_paragraph(
    "برای استفاده از مقادیر در Contextهای ری‌اکت استفاده می‌شود. "
    "Context برای به‌اشتراک‌گذاری داده‌ها بین کامپوننت‌های مختلف، بدون نیاز به ارسال پراپ‌ها در هر سطح، کاربرد دارد."
)
doc.add_paragraph(
    """```javascript
import React, { useContext, createContext } from 'react';

const MyContext = createContext();

function MyComponent() {
  const value = useContext(MyContext);

  return <div>{value}</div>;
}

function App() {
  return (
    <MyContext.Provider value="Hello, World!">
      <MyComponent />
    </MyContext.Provider>
  );
}
```""",
    style='Code'
)

# useReducer
doc.add_heading('useReducer', level=2)
doc.add_paragraph(
    "برای مدیریت وضعیت‌های پیچیده‌تر که شامل چندین اکشن و حالت مختلف است استفاده می‌شود. "
    "useReducer مشابه یک ماشین حالت (state machine) عمل می‌کند که با استفاده از اکشن‌ها وضعیت را تغییر می‌دهد. "
    "در مقایسه با useState، برای وضعیت‌های پیچیده‌تر و وابسته به هم مناسب‌تر است."
)
doc.add_paragraph(
    """```javascript
import React, { useReducer } from 'react';

const initialState = { count: 0 };

function reducer(state, action) {
  switch (action.type) {
    case 'increment':
      return { count: state.count + 1 };
    case 'decrement':
      return { count: state.count - 1 };
    default:
      throw new Error();
  }
}

function Counter() {
  const [state, dispatch] = useReducer(reducer, initialState);

  return (
    <div>
      <p>Count: {state.count}</p>
      <button onClick={() => dispatch({ type: 'increment' })}>Increment</button>
      <button onClick={() => dispatch({ type: 'decrement' })}>Decrement</button>
    </div>
  );
}
```""",
    style='Code'
)

# useMemo
doc.add_heading('useMemo', level=2)
doc.add_paragraph(
    "برای محاسبات پرهزینه که باید در رندر مجدد حفظ شوند، استفاده می‌شود. "
    "useMemo مقداری را فقط زمانی محاسبه می‌کند که یکی از وابستگی‌ها تغییر کند."
)
doc.add_paragraph(
    """```javascript
import React, { useMemo, useState } from 'react';

function ExpensiveComponent({ number }) {
  const computeExpensiveValue = (num) => {
    console.log('Computing...');
    // شبیه‌سازی محاسبات سنگین
    return num * 2;
  };

  const expensiveValue = useMemo(() => computeExpensiveValue(number), [number]);

  return <div>{expensiveValue}</div>;
}

function App() {
  const [count, setCount] = useState(0);

  return (
    <div>
      <ExpensiveComponent number={count} />
      <button onClick={() => setCount(count + 1)}>Increment</button>
    </div>
  );
}
```""",
    style='Code'
)

# useCallback
doc.add_heading('useCallback', level=2)
doc.add_paragraph(
    "برای حفظ مراجع توابع بین رندرها استفاده می‌شود، به خصوص زمانی که این توابع به عنوان پراپ به کامپوننت‌های فرزند ارسال می‌شوند."
)
doc.add_paragraph(
    """```javascript
import React, { useCallback, useState } from 'react';

function Child({ onClick }) {
  return <button onClick={onClick}>Click me</button>;
}

function Parent() {
  const [count, setCount] = useState(0);

  const handleClick = useCallback(() => {
    setCount(count + 1);
  }, [count]);

  return (
    <div>
      <Child onClick={handleClick} />
      <p>Count: {count}</p>
    </div>
  );
}
```""",
    style='Code'
)

# useRef
doc.add_heading('useRef', level=2)
doc.add_paragraph(
    "برای دسترسی به عناصر DOM و ذخیره مقادیر که باعث رندر مجدد نمی‌شوند، استفاده می‌شود."
)
doc.add_paragraph(
    """```javascript
import React, { useRef } from 'react';

function TextInputWithFocusButton() {
  const inputEl = useRef(null);

  const onButtonClick = () => {
    // با استفاده از useRef به عنصر DOM دسترسی پیدا می‌کنیم
    inputEl.current.focus();
  };

  return (
    <div>
      <input ref={inputEl} type="text" />
      <button onClick={onButtonClick}>Focus the input</button>
    </div>
  );
}
```""",
    style='Code'
)

# useLayoutEffect
doc.add_heading('useLayoutEffect', level=2)
doc.add_paragraph(
    "شبیه به useEffect، اما بعد از همه تغییرات DOM انجام می‌شود. "
    "به‌طور معمول برای اثرات جانبی که نیاز به اندازه‌گیری یا تغییرات DOM دارند استفاده می‌شود."
)
doc.add_paragraph(
    """```javascript
import React, { useLayoutEffect, useRef } from 'react';

function LayoutEffectExample() {
  const divRef = useRef(null);

  useLayoutEffect(() => {
    console.log(divRef.current.getBoundingClientRect());
  }, []);

  return <div ref={divRef}>Hello, World!</div>;
}
```""",
    style='Code'
)

# useImperativeHandle
doc.add_heading('useImperativeHandle', level=2)
doc.add_paragraph(
    "برای سفارشی‌سازی مقادیر قابل دسترس از طریق ref استفاده می‌شود. "
    "همراه با forwardRef برای ایجاد APIهای کامپوننت‌های سفارشی استفاده می‌شود."
)
doc.add_paragraph(
    """```javascript
import React, { useImperativeHandle, forwardRef, useRef } from 'react';

const FancyInput = forwardRef((props, ref) => {
  const inputRef = useRef();

  useImperativeHandle(ref, () => ({
    focus: () => {
      inputRef.current.focus();
    }
  }));

  return <input ref={inputRef} />;
});

function Parent() {
  const ref = useRef();

  return (
    <div>
      <FancyInput ref={ref} />
      <button onClick={() => ref.current.focus()}>Focus the input</button>
    </div>
  );
}
```""",
    style='Code'
)

# هوک‌های نکست‌جی‌اس
doc.add_heading('هوک‌های نکست‌جی‌اس', level=1)

# useRouter
doc.add_heading('useRouter', level=2)
doc.add_paragraph(
    "برای دسترسی به شیء مسیریاب (router) در کامپوننت‌های فانکشنال نکست‌جی‌اس استفاده می‌شود. "
    "این هوک امکاناتی مانند پیمایش برنامه‌ای (programmatic navigation) و دسترسی به پارامترهای مسیر را فراهم می‌کند."
)
doc.add_paragraph(
    """```javascript
import { useRouter } from 'next/router';

function MyComponent() {
  const router = useRouter();

  const goToHome = () => {
    router.push('/');
  };

  return <button onClick={goToHome}>Go to Home</button>;
}
```""",
    style='Code'
)

# useSWR
doc.add_heading('useSWR', level=2)
doc.add_paragraph(
    "برای مدیریت درخواست‌های داده و کش کردن نتایج استفاده می‌شود. "
    "SWR یک کتابخانه برای fetching داده‌ها است که توسط تیم ورسل توسعه داده شده است و با نکست‌جی‌اس به خوبی سازگار است."
)
doc.add_paragraph(
    """```javascript
import useSWR from 'swr';

function Fetcher() {
  const fetcher = (url) => fetch(url).then((res) => res.json());
  const { data, error } = useSWR('/api/data', fetcher);

  if (error) return <div>Failed to load</div>;
  if (!data) return <div>Loading...</div>;

  return <div>{data.message}</div>;
}
```""",
    style='Code'
)

# ذخیره سند
doc.save("React_NextJS_Hooks.docx")
